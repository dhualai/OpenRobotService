"""
车端界面中英文对照表 .docx → Qdrant

读取 cheduan_doc/车载界面中英文对照表.docx，分两部分入库：
- 表 0~15：UI 中英对照 → translation 集合
- 表 16~23：3 位错误码（200~916）→ cheduan 集合

使用方法：
    python -m ai.ingestion.ingest_cheduan_docx --rebuild
    python -m ai.ingestion.ingest_cheduan_docx --dry-run
    python -m ai.ingestion.ingest_cheduan_docx --cleanup
"""
import sys
import hashlib
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime

_project_root = Path(__file__).resolve().parent.parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from dotenv import load_dotenv
load_dotenv(_project_root / "ai" / ".env")

from ai.config import get_docs_dir

DOCS = get_docs_dir()
DOCX_PATH = DOCS / "cheduan_doc" / "车载界面中英文对照表.docx"
COLLECTION_PREFIX = "cheduan"


# ============================================================
# 数据模型
# ============================================================

@dataclass
class CheduanErrorEntry:
    """3 位车端错误码（来自 docx 表 16~23）"""
    module: str           # 模块名：运动控制/定位/Monitor/感知/底层驱动/电池/网关
    code: str             # 错误码（200~916）
    level: str            # 等级：警告/故障/提示
    description_cn: str   # 中文描述
    description_en: str   # 英文描述


@dataclass
class UITranslationEntry:
    """UI 标签中英对照（来自 docx 表 0~15）"""
    cn: str
    en: str
    category: str = ""   # 分类：主界面/车辆状态/VDA5050 等


# ============================================================
# 解析
# ============================================================

def _cell_text(row, col_idx: int) -> str:
    """安全取单元格文本"""
    if col_idx >= len(row.cells):
        return ""
    return (row.cells[col_idx].text or "").strip()


def parse_error_codes(docx_path: Path) -> List[CheduanErrorEntry]:
    """解析表 16~23：3 位车端错误码"""
    from docx import Document
    doc = Document(str(docx_path))
    entries: List[CheduanErrorEntry] = []
    seen: set = set()

    for ti in range(16, min(24, len(doc.tables))):
        table = doc.tables[ti]
        module = _cell_text(table.rows[0], 0) or "未知模块"

        for ri in range(1, len(table.rows)):
            row = table.rows[ri]
            code = _cell_text(row, 0)
            level = _cell_text(row, 1)
            desc_cn = _cell_text(row, 2)
            level_en = _cell_text(row, 3)
            desc_en = _cell_text(row, 4)

            # 过滤表头/空行
            if not code or not code.isdigit():
                continue
            if not desc_cn:
                continue

            # 等级优先用中文列，fallback 到英文列再翻译
            level_map = {
                "警告": "Warn", "故障": "Error", "提示": "Info",
                "Warning": "Warn", "Fault": "Error", "Prompt": "Info",
            }
            level_normalized = level_map.get(level) or level_map.get(level_en) or "Warn"

            key = f"{code}_{module}"
            if key in seen:
                continue
            seen.add(key)

            entries.append(CheduanErrorEntry(
                module=module,
                code=code,
                level=level_normalized,
                description_cn=desc_cn,
                description_en=desc_en,
            ))

    return entries


def parse_ui_translations(docx_path: Path) -> List[UITranslationEntry]:
    """解析表 0~15：UI 中英对照"""
    from docx import Document
    doc = Document(str(docx_path))
    entries: List[UITranslationEntry] = []
    current_category = ""

    for ti in range(0, 16):
        if ti >= len(doc.tables):
            break
        table = doc.tables[ti]
        if len(table.rows) < 2:
            continue

        # 从表之前的段落提取分类名
        header_text = _cell_text(table.rows[0], 0)

        for ri in range(1, len(table.rows)):
            row = table.rows[ri]
            cn = _cell_text(row, 0)
            en = _cell_text(row, 1)

            # 跳过表头行和空行
            if not cn or cn in ("中文", "English"):
                continue
            if cn == en and cn:
                # 可能是分类标题（如 "车辆状态"、"VDA5050"）
                if any('一' <= c <= '鿿' for c in cn):
                    current_category = cn
                continue
            if not en:
                continue

            entries.append(UITranslationEntry(
                cn=cn,
                en=en,
                category=current_category or header_text,
            ))

    return entries


# ============================================================
# Chunk 构建
# ============================================================

def build_error_chunks(entries: List[CheduanErrorEntry]) -> List[Dict[str, Any]]:
    """错误码 → cheduan collection 格式"""
    chunks = []
    for e in entries:
        text = (
            f"【车端错误码】{e.code}\n"
            f"模块：{e.module}\n"
            f"类别：车端本体错误\n"
            f"等级：{e.level}\n"
            f"描述：{e.description_cn}\n"
            f"Description: {e.description_en}"
        )
        chunks.append({
            "id": hashlib.md5(f"cheduan_v2_{e.code}".encode()).hexdigest(),
            "text": text,
            "payload": {
                "error_code": e.code,
                "category": e.module,
                "level": e.level,
                "description_cn": e.description_cn,
                "description_en": e.description_en,
                "solution_cn": "",
                "solution_en": "",
                "source": "车载界面中英文对照表.docx",
            },
        })
    return chunks


def build_translation_chunks(entries: List[UITranslationEntry]) -> List[Dict[str, Any]]:
    """UI 翻译 → translation collection 格式"""
    chunks = []
    for e in entries:
        text = f"{e.cn} | {e.en}"
        chunks.append({
            "id": hashlib.md5(f"ui_trans_{e.cn}_{e.en}".encode()).hexdigest(),
            "text": text,
            "payload": {
                "cn": e.cn,
                "en": e.en,
                "category": e.category,
                "source": "车载界面中英文对照表.docx",
            },
        })
    return chunks


# ============================================================
# 通用入库
# ============================================================

async def _ingest_chunks(
    chunks: List[Dict],
    collection_name: str,
    label: str,
    rebuild: bool = False,
) -> Optional[Dict]:
    from ai.core.embed import get_embed_client
    from ai.config import get_ai_config
    from qdrant_client import QdrantClient
    from qdrant_client.models import PointStruct, Distance, VectorParams

    config = get_ai_config()
    embed_client = await get_embed_client()

    print(f"\n[INGEST] {len(chunks)} {label} chunks -> collection: {collection_name}")

    texts = [c["text"] for c in chunks]
    vectors = await embed_client.embed_batch(texts)
    dim = vectors[0].shape[-1]

    if config.qdrant_local_path:
        local = Path(config.qdrant_local_path)
        if not local.is_absolute():
            local = _project_root / local
        client = QdrantClient(path=str(local))
    else:
        client = QdrantClient(
            host=config.qdrant_host, port=config.qdrant_port,
            timeout=config.qdrant_timeout, check_compatibility=False,
        )

    if rebuild and client.collection_exists(collection_name):
        try:
            client.delete_collection(collection_name)
        except Exception:
            pass

    if not client.collection_exists(collection_name):
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
        )

    points = [
        PointStruct(id=c["id"], vector=v.tolist(), payload=c["payload"])
        for c, v in zip(chunks, vectors)
    ]
    for i in range(0, len(points), 50):
        client.upsert(collection_name=collection_name, points=points[i:i+50])

    return {"status": "ok", "entries": len(chunks), "dimension": dim, "collection": collection_name}


async def _cleanup_old(prefix: str, keep: int = 1):
    from ai.config import get_ai_config
    from qdrant_client import QdrantClient

    config = get_ai_config()
    if config.qdrant_local_path:
        local = Path(config.qdrant_local_path)
        if not local.is_absolute():
            local = _project_root / local
        client = QdrantClient(path=str(local))
    else:
        client = QdrantClient(
            host=config.qdrant_host, port=config.qdrant_port,
            timeout=config.qdrant_timeout, check_compatibility=False,
        )

    try:
        all_cols = [c.name for c in client.get_collections().collections]
    except Exception as e:
        print(f"[ERR] 获取集合列表失败: {e}")
        return

    ours = sorted([c for c in all_cols if c.startswith(prefix)], reverse=True)
    for c in ours[keep:]:
        try:
            client.delete_collection(c)
            print(f"  [DEL] {c}")
        except Exception:
            pass


# ============================================================
# 自动入库入口
# ============================================================

async def auto_ingest() -> bool:
    """启动时自动调用"""
    from ai.config import _write_active_cheduan_collection, _write_active_translation_collection
    from ai.config import get_active_cheduan_collection, get_active_translation_collection

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    # ── 错误码（3 位）→ cheduan 集合 ──
    error_entries = parse_error_codes(DOCX_PATH)
    print(f"[CHE DUAN DOCX] 解析到 {len(error_entries)} 条 3 位错误码")

    if error_entries:
        error_chunks = build_error_chunks(error_entries)
        # 追加到现有 cheduan 集合（保留 PDF 中的 4-5 位码）
        cheduan_col = get_active_cheduan_collection()
        if not cheduan_col:
            cheduan_col = f"{COLLECTION_PREFIX}_{ts}"

        await _ingest_chunks(error_chunks, cheduan_col, "车端3位错误码", rebuild=False)

        old_cd = get_active_cheduan_collection()
        if old_cd != cheduan_col:
            _write_active_cheduan_collection(cheduan_col)
        print(f"[SWITCH] cheduan: {old_cd or '(new)'} -> {cheduan_col}")
    else:
        print("[CHE DUAN DOCX] 未解析到错误码")

    # ── UI 翻译 → translation 集合 ──
    ui_entries = parse_ui_translations(DOCX_PATH)
    print(f"[CHE DUAN DOCX] 解析到 {len(ui_entries)} 条 UI 翻译")

    if ui_entries:
        ui_chunks = build_translation_chunks(ui_entries)
        translation_col_name = f"translation_v2_{ts}"

        await _ingest_chunks(ui_chunks, translation_col_name, "车端UI翻译", rebuild=True)

        old_tr = get_active_translation_collection()
        _write_active_translation_collection(translation_col_name)
        print(f"[SWITCH] translation: {old_tr or '(new)'} -> {translation_col_name}")
        await _cleanup_old("translation", keep=2)
    else:
        print("[CHE DUAN DOCX] 未解析到 UI 翻译")

    return True


# ============================================================
# CLI
# ============================================================

async def main():
    import argparse
    from ai.config import get_active_cheduan_collection, get_active_translation_collection
    from ai.config import _write_active_cheduan_collection, _write_active_translation_collection

    parser = argparse.ArgumentParser(description="车端界面中英文对照表 .docx -> Qdrant")
    parser.add_argument("--rebuild", "-r", action="store_true", help="完整入库")
    parser.add_argument("--dry-run", "-n", action="store_true", help="预览")
    parser.add_argument("--cleanup", action="store_true", help="清理旧集合")
    args = parser.parse_args()

    if args.cleanup:
        await _cleanup_old(COLLECTION_PREFIX, keep=2)
        await _cleanup_old("translation", keep=2)
        return

    # ── 错误码 ──
    error_entries = parse_error_codes(DOCX_PATH)
    print(f"[CHE DUAN DOCX] 解析到 {len(error_entries)} 条 3 位错误码")

    if args.dry_run:
        print("\n错误码预览（前 10 条）:\n")
        for e in error_entries[:10]:
            print(f"  [{e.level}] {e.code} | {e.module} | {e.description_cn[:60]}")
        print("\nUI 翻译预览（前 10 条）:\n")
        ui_entries = parse_ui_translations(DOCX_PATH)
        for e in ui_entries[:10]:
            print(f"  [{e.category}] {e.cn} | {e.en}")
        return

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    if error_entries:
        error_chunks = build_error_chunks(error_entries)
        cheduan_col = get_active_cheduan_collection()
        if not cheduan_col:
            cheduan_col = f"{COLLECTION_PREFIX}_{ts}"

        result = await _ingest_chunks(error_chunks, cheduan_col, "车端3位错误码", rebuild=False)
        if result and result["status"] == "ok":
            old_cd = get_active_cheduan_collection()
            if old_cd != cheduan_col:
                _write_active_cheduan_collection(cheduan_col)
            print(f"[SWITCH] cheduan: {old_cd or '(new)'} -> {cheduan_col}")

    # ── UI 翻译 ──
    ui_entries = parse_ui_translations(DOCX_PATH)
    print(f"[CHE DUAN DOCX] 解析到 {len(ui_entries)} 条 UI 翻译")

    if ui_entries:
        ui_chunks = build_translation_chunks(ui_entries)
        translation_col_name = f"translation_v2_{ts}"

        old_tr = get_active_translation_collection()
        result = await _ingest_chunks(ui_chunks, translation_col_name, "车端UI翻译", rebuild=args.rebuild)
        if result and result["status"] == "ok":
            _write_active_translation_collection(translation_col_name)
            print(f"[SWITCH] translation: {old_tr or '(new)'} -> {translation_col_name}")
            await _cleanup_old("translation", keep=2)


if __name__ == "__main__":
    asyncio.run(main())
