"""
车端界面中英文对照表 .docx → cheduan 集合 (3 位错误码)

来源：cheduan_doc/车载界面中英文对照表.docx → 表 16~23
产出：120 条 3 位错误码（200~916）→ 追加到现有 cheduan 集合
"""
from pathlib import Path
from typing import List
from dataclasses import dataclass

from ai.config import get_docs_dir
from ai.ingestion.base import BaseIngester, Chunk
from ai.ingestion.registry import register


@dataclass
class CheduanDocxErrorEntry:
    """3 位车端错误码（来自 docx 表 16~23）"""
    module: str
    code: str
    level: str            # Warn / Error / Info
    description_cn: str
    description_en: str


class CheduanDocxIngester(BaseIngester[CheduanDocxErrorEntry]):
    """车端 docx 3位错误码（200~916）→ cheduan 集合（追加模式）"""

    source_paths = [get_docs_dir() / "cheduan_doc" / "车载界面中英文对照表.docx"]
    collection_prefix = "cheduan"
    collection_type = "cheduan"
    rebuild = False  # 追加到现有集合（保留 PDF 4-5 位码）

    @staticmethod
    def _pointer_reader() -> str:
        from ai.config import get_active_cheduan_collection
        return get_active_cheduan_collection()

    @staticmethod
    def _pointer_writer(name: str) -> None:
        from ai.config import _write_active_cheduan_collection
        _write_active_cheduan_collection(name)

    pointer_reader = staticmethod(_pointer_reader)
    pointer_writer = staticmethod(_pointer_writer)

    def parse(self) -> List[CheduanDocxErrorEntry]:
        return _parse_error_tables(self.source_paths[0])

    def to_chunk(self, e: CheduanDocxErrorEntry) -> Chunk:
        text = (
            f"【车端错误码】{e.code}\n"
            f"模块：{e.module}\n"
            f"类别：车端本体错误\n"
            f"等级：{e.level}\n"
            f"描述：{e.description_cn}\n"
            f"Description: {e.description_en}"
        )
        return Chunk(
            id=self.stable_id("cheduan", "v2", e.code),
            text=text,
            payload={
                "error_code": e.code,
                "category": e.module,
                "level": e.level,
                "description_cn": e.description_cn,
                "description_en": e.description_en,
                "solution_cn": "",
                "solution_en": "",
                "source": "车载界面中英文对照表.docx",
            },
        )


def _parse_error_tables(docx_path: Path) -> List[CheduanDocxErrorEntry]:
    """解析表 16~23：3 位车端错误码"""
    from docx import Document

    doc = Document(str(docx_path))
    entries: List[CheduanDocxErrorEntry] = []
    seen: set = set()

    for ti in range(16, min(24, len(doc.tables))):
        table = doc.tables[ti]
        module = _cell_text(table.rows[0], 0) or "未知模块"

        for ri in range(1, len(table.rows)):
            row = table.rows[ri]
            code = _cell_text(row, 0)
            level = _cell_text(row, 1)
            desc_cn = _cell_text(row, 2)
            level_en = _cell_text(row, 3)
            desc_en = _cell_text(row, 4)

            if not code or not code.isdigit():
                continue
            if not desc_cn:
                continue

            level_map = {
                "警告": "Warn", "故障": "Error", "提示": "Info",
                "Warning": "Warn", "Fault": "Error", "Prompt": "Info",
            }
            level_normalized = level_map.get(level) or level_map.get(level_en) or "Warn"

            key = f"{code}_{module}"
            if key in seen:
                continue
            seen.add(key)

            entries.append(CheduanDocxErrorEntry(
                module=module,
                code=code,
                level=level_normalized,
                description_cn=desc_cn,
                description_en=desc_en,
            ))

    return entries


def _cell_text(row, col_idx: int) -> str:
    """安全取单元格文本"""
    if col_idx >= len(row.cells):
        return ""
    return (row.cells[col_idx].text or "").strip()


def register_all():
    register(CheduanDocxIngester, description="车端 docx 3位错误码 → cheduan 集合（追加）")


register_all()
