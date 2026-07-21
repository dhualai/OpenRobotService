"""
车端 UI 中英对照 .docx → translation 集合

来源：cheduan_doc/车载界面中英文对照表.docx → 表 0~15
产出：192 条 UI 翻译 → translation 集合（独立 collection）
"""
from pathlib import Path
from typing import List
from dataclasses import dataclass

from ai.config import get_docs_dir
from ai.ingestion.base import BaseIngester, Chunk
from ai.ingestion.registry import register


@dataclass
class UITranslationEntry:
    cn: str
    en: str
    category: str = ""


class UITranslationDocxIngester(BaseIngester[UITranslationEntry]):
    """车端 docx UI 中英对照（表 0~15）→ translation 集合"""

    source_paths = [get_docs_dir() / "cheduan_doc" / "车载界面中英文对照表.docx"]
    collection_prefix = "translation"
    collection_type = "translation"
    rebuild = True  # 独立 collection，完全重建

    @staticmethod
    def _pointer_reader() -> str:
        from ai.config import get_active_translation_collection
        return get_active_translation_collection()

    @staticmethod
    def _pointer_writer(name: str) -> None:
        from ai.config import _write_active_translation_collection
        _write_active_translation_collection(name)

    pointer_reader = staticmethod(_pointer_reader)
    pointer_writer = staticmethod(_pointer_writer)

    def parse(self) -> List[UITranslationEntry]:
        return _parse_ui_tables(self.source_paths[0])

    def to_chunk(self, e: UITranslationEntry) -> Chunk:
        text = f"{e.cn} | {e.en}"
        return Chunk(
            id=self.stable_id("ui_trans", e.cn, e.en),
            text=text,
            payload={
                "cn": e.cn,
                "en": e.en,
                "category": e.category,
                "source": "车载界面中英文对照表.docx",
            },
        )


def _parse_ui_tables(docx_path: Path) -> List[UITranslationEntry]:
    """解析表 0~15：UI 中英对照"""
    from docx import Document

    doc = Document(str(docx_path))
    entries: List[UITranslationEntry] = []
    current_category = ""

    for ti in range(0, 16):
        if ti >= len(doc.tables):
            break
        table = doc.tables[ti]
        if len(table.rows) < 2:
            continue

        header_text = _cell_text(table.rows[0], 0)

        for ri in range(1, len(table.rows)):
            row = table.rows[ri]
            cn = _cell_text(row, 0)
            en = _cell_text(row, 1)

            if not cn or cn in ("中文", "English"):
                continue
            if cn == en and cn:
                if any('一' <= c <= '鿿' for c in cn):
                    current_category = cn
                continue
            if not en:
                continue

            entries.append(UITranslationEntry(
                cn=cn,
                en=en,
                category=current_category or header_text,
            ))

    return entries


def _cell_text(row, col_idx: int) -> str:
    if col_idx >= len(row.cells):
        return ""
    return (row.cells[col_idx].text or "").strip()


def register_all():
    register(UITranslationDocxIngester, description="车端 docx UI 中英对照（表0~15）→ translation 集合")


register_all()
